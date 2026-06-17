"""Markdown → DOCX exporter.

Lightweight markdown subset (good enough for synth-emitted briefs):
  - `# h1` / `## h2` / `### h3` headings
  - `- item` / `* item` bullet lists
  - blank-line-separated paragraphs
  - **bold** / *italic* inline runs

`build_docx(title, body_md, out_path)` writes a .docx and returns the
absolute path. Never raises on parse weirdness — falls back to plain
paragraphs.
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.shared import Pt


_HEADING_RE = re.compile(r"^(#{1,3})\s+(.+?)\s*$")
_BULLET_RE = re.compile(r"^[\-\*]\s+(.+?)\s*$")
_INLINE_RUN_RE = re.compile(
    r"(\*\*(?P<bold>[^*]+)\*\*|\*(?P<italic>[^*]+)\*)"
)


def _add_inline_runs(paragraph, text: str) -> None:
    """Render **bold** and *italic* spans as separate runs."""
    pos = 0
    for m in _INLINE_RUN_RE.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        if m.group("bold") is not None:
            r = paragraph.add_run(m.group("bold"))
            r.bold = True
        elif m.group("italic") is not None:
            r = paragraph.add_run(m.group("italic"))
            r.italic = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def build_docx(title: str, body_md: str, out_path: Path) -> Path:
    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    # Body default — 11pt, mirrors what most readers expect.
    style = doc.styles["Normal"]
    style.font.size = Pt(11)

    if title:
        doc.add_heading(title, level=0)

    lines = (body_md or "").splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line.strip():
            i += 1
            continue

        m = _HEADING_RE.match(line)
        if m:
            doc.add_heading(m.group(2).strip(), level=len(m.group(1)))
            i += 1
            continue

        m = _BULLET_RE.match(line)
        if m:
            while i < len(lines):
                bm = _BULLET_RE.match(lines[i].rstrip())
                if not bm:
                    break
                p = doc.add_paragraph(style="List Bullet")
                _add_inline_runs(p, bm.group(1))
                i += 1
            continue

        # Paragraph: collect until blank line.
        para_lines: list[str] = [line]
        i += 1
        while i < len(lines) and lines[i].strip():
            if _HEADING_RE.match(lines[i]) or _BULLET_RE.match(lines[i]):
                break
            para_lines.append(lines[i].rstrip())
            i += 1
        p = doc.add_paragraph()
        _add_inline_runs(p, " ".join(para_lines))

    doc.save(str(out_path))
    return out_path.resolve()


__all__ = ["build_docx"]
